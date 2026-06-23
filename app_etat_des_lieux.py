#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Application Streamlit - Etat des lieux Lycee
Lancer avec : streamlit run app_etat_des_lieux.py
"""

import io
import unicodedata
from datetime import date
from pathlib import Path

import pandas as pd
import streamlit as st

# ── Config page ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Etat des lieux - Lycee",
    page_icon="🏫",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Constantes ─────────────────────────────────────────────────────────────────
VALEURS_KO = {
    "manquant":          ("CRITIQUE",     "🔴"),
    "manquante":         ("CRITIQUE",     "🔴"),
    "a changer":         ("REMPLACEMENT", "🟠"),
    "travaux a prevoir": ("TRAVAUX",      "🟡"),
    "dalles":            ("TRAVAUX",      "🟡"),
    "changer piles":     ("REMPLACEMENT", "🟠"),
    "deteriore":         ("CRITIQUE",     "🔴"),
    "deterioree":        ("CRITIQUE",     "🔴"),
}
VALEURS_OK = {"bon", "tres bon", "presente", "present", "x", "", "none", "nan"}

COULEURS_PRIO = {
    "CRITIQUE":     "#C0392B",  # 5,41:1 sur blanc, 4,78:1 sur fond rose  (RGAA ✅)
    "REMPLACEMENT": "#A04000",  # 6,55:1 sur blanc, 5,99:1 sur fond orange (RGAA ✅)
    "TRAVAUX":      "#7D6608",  # 5,59:1 sur blanc, 5,44:1 sur fond jaune  (RGAA ✅)
    "A VERIFIER":   "#5D5D5D",  # 6,86:1 sur blanc                         (RGAA ✅)
}
FONDS_PRIO = {
    "CRITIQUE":     "#FFEBEB",
    "REMPLACEMENT": "#FFF3E0",
    "TRAVAUX":      "#FFFDE7",
    "A VERIFIER":   "#F5F5F5",
}
ORDRE_PRIO = ["CRITIQUE", "REMPLACEMENT", "TRAVAUX", "A VERIFIER"]

# ── Utilitaires ────────────────────────────────────────────────────────────────
def normaliser(s: str) -> str:
    return (
        unicodedata.normalize("NFD", str(s))
        .encode("ascii", "ignore")
        .decode("ascii")
        .lower()
        .strip()
    )


def determiner_priorite(valeur: str):
    norm = normaliser(valeur)
    if norm in VALEURS_OK:
        return None
    matched = VALEURS_KO.get(norm)
    if matched is None:
        for motcle, info in VALEURS_KO.items():
            if motcle in norm:
                matched = info
                break
    if matched:
        return matched  # (priorite, emoji)
    return ("A VERIFIER", "⚪")


@st.cache_data
def charger_xlsx(contenu: bytes) -> pd.DataFrame:
    """Charge le xlsx et retourne un DataFrame plat des interventions."""
    xl = pd.ExcelFile(io.BytesIO(contenu))
    ws = xl.parse(xl.sheet_names[0], header=1)
    ws.columns = [str(c).strip() for c in ws.columns]
    col_salle = ws.columns[0]

    rows = []
    for _, row in ws.iterrows():
        salle = str(row[col_salle]).strip()
        if salle in ("nan", "None", ""):
            continue
        for col in ws.columns[1:]:
            if str(col).startswith("Unnamed") or str(col) == "nan":
                continue
            valeur = str(row[col]).strip() if pd.notna(row[col]) else ""
            res = determiner_priorite(valeur)
            if res:
                prio, emoji = res
                rows.append({
                    "Salle":      salle,
                    "Element":    col,
                    "Valeur":     valeur,
                    "Priorite":   prio,
                    "Emoji":      emoji,
                    "Traite":     False,
                })
    return pd.DataFrame(rows)


# ── Export DOCX ────────────────────────────────────────────────────────────────
def generer_docx(df: pd.DataFrame) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def hex2rgb(h):
        h = h.lstrip("#")
        return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))

    def set_cell_bg(cell, hex_color):
        hex_color = hex_color.lstrip("#")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_para_bg(para, hex_color):
        hex_color = hex_color.lstrip("#")
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def add_border_bottom(para, color="1A3C6E"):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "1");    b.set(qn("w:color"), color)
        pBdr.append(b); pPr.append(pBdr)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2); section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

    # Titre
    p = doc.add_paragraph()
    r = p.add_run("Etat des lieux - Interventions a entreprendre")
    r.bold = True; r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(p)

    p2 = doc.add_paragraph()
    p2.add_run(f"Lycee - Annee 2025/2026  |  Genere le {date.today().strftime('%d/%m/%Y')}")
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Bilan global
    total_i = len(df)
    non_traites = len(df[~df["Traite"]]) if "Traite" in df.columns else total_i
    p_stat = doc.add_paragraph()
    r_stat = p_stat.add_run(
        f"Total interventions : {total_i}  |  Non traitees : {non_traites}"
    )
    r_stat.bold = True; r_stat.font.size = Pt(11)
    r_stat.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p_stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    for salle in df["Salle"].unique():
        df_s = df[df["Salle"] == salle]
        # En-tete salle
        p_hdr = doc.add_paragraph()
        r_hdr = p_hdr.add_run(f"  Salle : {salle}  ({len(df_s)} intervention(s))")
        r_hdr.bold = True; r_hdr.font.size = Pt(12)
        r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_para_bg(p_hdr, "1A3C6E")

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells,
                               ["Element", "Valeur originale", "Priorite", "Statut"]):
            cell.text = label
            set_cell_bg(cell, "D5E8F0")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for _, row in df_s.sort_values("Priorite").iterrows():
            cells = table.add_row().cells
            cells[0].text = str(row["Element"])
            cells[1].text = str(row["Valeur"])
            cells[2].text = str(row["Priorite"])
            cells[3].text = "✓ Traite" if row.get("Traite") else "En attente"

            bg = FONDS_PRIO.get(row["Priorite"], "#F5F5F5").lstrip("#")
            tc_hex = COULEURS_PRIO.get(row["Priorite"], "#888888").lstrip("#")
            set_cell_bg(cells[2], bg)
            r_prio = cells[2].paragraphs[0].runs[0]
            r_prio.bold = True; r_prio.font.size = Pt(9)
            r_prio.font.color.rgb = RGBColor(*hex2rgb(tc_hex))
            for c in cells:
                c.paragraphs[0].runs[0].font.size = Pt(9)

        for row in table.rows:
            row.cells[0].width = Cm(5)
            row.cells[1].width = Cm(5)
            row.cells[2].width = Cm(3.5)
            row.cells[3].width = Cm(2.5)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Export XLSX ────────────────────────────────────────────────────────────────
def generer_xlsx(df: pd.DataFrame) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        export_df = df[["Salle", "Element", "Valeur", "Priorite", "Traite"]].copy()
        export_df["Traite"] = export_df["Traite"].map({True: "Oui", False: "Non"})
        export_df.to_excel(writer, index=False, sheet_name="Interventions")

        ws = writer.sheets["Interventions"]
        from openpyxl.styles import PatternFill, Font, Alignment
        FILLS = {
            "CRITIQUE":     PatternFill("solid", fgColor="FFEBEB"),
            "REMPLACEMENT": PatternFill("solid", fgColor="FFF3E0"),
            "TRAVAUX":      PatternFill("solid", fgColor="FFFDE7"),
            "A VERIFIER":   PatternFill("solid", fgColor="F5F5F5"),
        }
        FONTS = {
            "CRITIQUE":     Font(bold=True, color="C0392B"),
            "REMPLACEMENT": Font(bold=True, color="A04000"),
            "TRAVAUX":      Font(bold=True, color="7D6608"),
            "A VERIFIER":   Font(bold=True, color="5D5D5D"),
        }
        header_fill = PatternFill("solid", fgColor="1A3C6E")
        header_font = Font(bold=True, color="FFFFFF")
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

        for row in ws.iter_rows(min_row=2):
            prio = row[3].value
            if prio in FILLS:
                row[3].fill = FILLS[prio]
                row[3].font = FONTS[prio]

        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 25
        ws.column_dimensions["D"].width = 18
        ws.column_dimensions["E"].width = 10

    return buf.getvalue()


# ══════════════════════════════════════════════════════════════════════════════
# INTERFACE
# ══════════════════════════════════════════════════════════════════════════════

# ── CSS personnalise ───────────────────────────────────────────────────────────
st.markdown("""
<style>
    .main-title {
        font-size: 2rem; font-weight: 800; color: #1A3C6E;
        border-bottom: 3px solid #1A3C6E; padding-bottom: 0.3rem;
        margin-bottom: 1rem;
    }
    .kpi-card {
        background: white; border-radius: 12px; padding: 1rem 1.2rem;
        box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        border-left: 5px solid; text-align: center;
    }
    .kpi-num { font-size: 2.2rem; font-weight: 800; line-height: 1; }
    .kpi-label { font-size: 0.875rem; color: #555; margin-top: 4px; }
    .salle-header {
        background: #1A3C6E; color: white; padding: 0.4rem 0.8rem;
        border-radius: 6px; font-weight: 700; margin: 0.8rem 0 0.3rem;
    }
    .badge {
        display: inline-block; padding: 2px 10px; border-radius: 20px;
        font-size: 0.75rem; font-weight: 700;
    }
    div[data-testid="stSidebarContent"] { background: #EBF0F8; }
</style>
""", unsafe_allow_html=True)

# ── Sidebar ────────────────────────────────────────────────────────────────────
with st.sidebar:
    col_logo1, col_logo2 = st.columns(2)
    with col_logo1:
        st.image("logo.png", use_container_width=True)
    with col_logo2:
        st.image("LOL.jpg", use_container_width=True)
    st.markdown("## 🏫 Etat des lieux Lycee")
    st.divider()

    fichier = st.file_uploader(
        "📂 Charger le fichier Excel",
        type=["xlsx"],
        help="Fichier Etat_des_lieux_2025_2026.xlsx",
    )

    if fichier:
        st.success(f"✅ {fichier.name}")
        st.divider()

        # Charger et stocker en session
        if "df_base" not in st.session_state or st.session_state.get("fichier_nom") != fichier.name:
            st.session_state.df_base = charger_xlsx(fichier.read())
            st.session_state.fichier_nom = fichier.name
            # Colonne Traite editable
            st.session_state.traites = {i: False for i in st.session_state.df_base.index}

        df_base = st.session_state.df_base.copy()
        df_base["Traite"] = df_base.index.map(st.session_state.traites)

        # ── Filtres ─────────────────────────────────────────────────────────
        st.markdown("### 🔍 Filtres")

        salles_dispo = sorted(df_base["Salle"].unique())
        salles_sel = st.multiselect(
            "Salle(s)", salles_dispo,
            placeholder="Toutes les salles",
        )

        prios_sel = st.multiselect(
            "Priorite(s)", ORDRE_PRIO,
            default=ORDRE_PRIO,
            help="Filtrer par niveau de priorite",
        )

        elements_dispo = sorted(df_base["Element"].unique())
        elements_sel = st.multiselect(
            "Element(s)", elements_dispo,
            placeholder="Tous les elements",
        )

        afficher_traites = st.toggle("Afficher les traitees", value=True)

        st.divider()
        st.markdown("### 📥 Exports")

# ── Corps principal ────────────────────────────────────────────────────────────
if not fichier:
    st.markdown('<div class="main-title">🏫 Etat des lieux - Lycee</div>', unsafe_allow_html=True)
    st.info(
        "👈 **Commencez par charger votre fichier Excel** dans la barre laterale.",
        icon="📂",
    )
    st.markdown("""
    #### Fonctionnalites disponibles
    - 📊 **Tableau de bord** avec statistiques globales
    - 🔍 **Filtres** par salle, priorite, element
    - ✅ **Marquer les interventions** comme traitees
    - 📄 **Export Word** (.docx) du rapport complet
    - 📊 **Export Excel** (.xlsx) du tableau filtre
    """)
    st.stop()

# ── Application des filtres ────────────────────────────────────────────────────
df = df_base.copy()
if salles_sel:
    df = df[df["Salle"].isin(salles_sel)]
if prios_sel:
    df = df[df["Priorite"].isin(prios_sel)]
if elements_sel:
    df = df[df["Element"].isin(elements_sel)]
if not afficher_traites:
    df = df[~df["Traite"]]

# ── Onglets ────────────────────────────────────────────────────────────────────
tab_dash, tab_liste, tab_salle, tab_edition = st.tabs([
    "📊 Tableau de bord",
    "📋 Liste des interventions",
    "🏠 Vue par salle",
    "✅ Gestion des traitements",
])


# ════════════════════════════════════════════════════════════════════════════
# TAB 1 : DASHBOARD
# ════════════════════════════════════════════════════════════════════════════
with tab_dash:
    st.markdown('<div class="main-title">📊 Tableau de bord</div>', unsafe_allow_html=True)

    df_tout = df_base.copy()
    total    = len(df_tout)
    critiques  = len(df_tout[df_tout["Priorite"] == "CRITIQUE"])
    remplace   = len(df_tout[df_tout["Priorite"] == "REMPLACEMENT"])
    travaux    = len(df_tout[df_tout["Priorite"] == "TRAVAUX"])
    traites    = len(df_tout[df_tout["Traite"]])
    nb_salles  = df_tout["Salle"].nunique()

    # KPI cards
    c1, c2, c3, c4, c5, c6 = st.columns(6)
    kpis = [
        (c1, total,     "Total",          "#1A3C6E"),
        (c2, critiques, "Critiques 🔴",   "#CC0000"),
        (c3, remplace,  "Remplacements 🟠","#A04000"),
        (c4, travaux,   "Travaux 🟡",     "#7D6608"),
        (c5, traites,   "Traitees ✅",    "#1E8449"),
        (c6, nb_salles, "Salles",         "#1A3C6E"),
    ]
    for col, val, label, color in kpis:
        with col:
            st.markdown(f"""
            <div class="kpi-card" style="border-color:{color}">
                <div class="kpi-num" style="color:{color}">{val}</div>
                <div class="kpi-label">{label}</div>
            </div>""", unsafe_allow_html=True)

    st.markdown("")

    # Graphiques
    col_g1, col_g2 = st.columns(2)

    with col_g1:
        st.markdown("#### Repartition par priorite")
        try:
            import plotly.express as px
            prio_counts = df_tout["Priorite"].value_counts().reindex(ORDRE_PRIO, fill_value=0)
            fig = px.pie(
                values=prio_counts.values,
                names=prio_counts.index,
                color=prio_counts.index,
                color_discrete_map={
                    "CRITIQUE": "#FF4444",
                    "REMPLACEMENT": "#E67E22",
                    "TRAVAUX": "#F0A500",
                    "A VERIFIER": "#888888",
                },
                hole=0.4,
            )
            fig.update_layout(margin=dict(t=10, b=10), height=300)
            st.plotly_chart(fig, use_container_width=True)
        except ImportError:
            st.bar_chart(df_tout["Priorite"].value_counts())

    with col_g2:
        st.markdown("#### Top 10 salles avec le plus d'interventions")
        try:
            top_salles = (
                df_tout[~df_tout["Traite"]]
                .groupby("Salle")
                .size()
                .sort_values(ascending=False)
                .head(10)
                .reset_index(name="nb")
            )
            fig2 = px.bar(
                top_salles, x="nb", y="Salle", orientation="h",
                color="nb",
                color_continuous_scale=["#FFFDE7", "#E67E22", "#CC0000"],
                labels={"nb": "Nb interventions", "Salle": ""},
            )
            fig2.update_layout(
                margin=dict(t=10, b=10), height=300,
                coloraxis_showscale=False,
                yaxis=dict(autorange="reversed"),
            )
            st.plotly_chart(fig2, use_container_width=True)
        except ImportError:
            st.bar_chart(
                df_tout.groupby("Salle").size().sort_values(ascending=False).head(10)
            )

    # Repartition par element
    st.markdown("#### Elements les plus souvent en cause")
    try:
        top_elem = (
            df_tout.groupby("Element").size()
            .sort_values(ascending=False).head(12)
            .reset_index(name="nb")
        )
        fig3 = px.bar(
            top_elem, x="Element", y="nb",
            color="nb",
            color_continuous_scale=["#FFFDE7", "#E67E22", "#CC0000"],
            labels={"nb": "Nb interventions", "Element": ""},
        )
        fig3.update_layout(margin=dict(t=10, b=10), height=280, coloraxis_showscale=False)
        st.plotly_chart(fig3, use_container_width=True)
    except ImportError:
        st.bar_chart(df_tout.groupby("Element").size().sort_values(ascending=False).head(12))


# ════════════════════════════════════════════════════════════════════════════
# TAB 2 : LISTE DES INTERVENTIONS
# ════════════════════════════════════════════════════════════════════════════
with tab_liste:
    st.markdown(f'<div class="main-title">📋 Interventions ({len(df)})</div>', unsafe_allow_html=True)

    if df.empty:
        st.success("✅ Aucune intervention selon les filtres appliques.")
    else:
        # Tri
        col_tri, col_ordre = st.columns([2, 1])
        with col_tri:
            tri = st.selectbox("Trier par", ["Priorite", "Salle", "Element", "Traite"])
        with col_ordre:
            ordre_asc = st.toggle("Croissant", value=True)

        ordre_prio_map = {"CRITIQUE": 0, "REMPLACEMENT": 1, "TRAVAUX": 2, "A VERIFIER": 3}
        df_affiche = df.copy()
        if tri == "Priorite":
            df_affiche["_ordre"] = df_affiche["Priorite"].map(ordre_prio_map)
            df_affiche = df_affiche.sort_values("_ordre", ascending=ordre_asc).drop(columns="_ordre")
        else:
            df_affiche = df_affiche.sort_values(tri, ascending=ordre_asc)

        # Affichage tableau
        df_display = df_affiche[["Emoji", "Salle", "Element", "Valeur", "Priorite", "Traite"]].copy()
        df_display.columns = ["", "Salle", "Element", "Valeur", "Priorite", "Traite"]
        df_display["Traite"] = df_display["Traite"].map({True: "✅ Oui", False: "⏳ Non"})

        def coloriser_ligne(row):
            bg = FONDS_PRIO.get(row["Priorite"], "#FFFFFF")
            return [f"background-color: {bg}"] * len(row)

        styled = df_display.style.apply(coloriser_ligne, axis=1)
        st.dataframe(styled, use_container_width=True, height=500, hide_index=True)


# ════════════════════════════════════════════════════════════════════════════
# TAB 3 : VUE PAR SALLE
# ════════════════════════════════════════════════════════════════════════════
with tab_salle:
    st.markdown('<div class="main-title">🏠 Vue par salle</div>', unsafe_allow_html=True)

    salles_filtrees = df["Salle"].unique() if not salles_sel else salles_sel

    if df.empty:
        st.success("✅ Aucune intervention pour les filtres selectionnes.")
    else:
        for salle in sorted(salles_filtrees):
            df_s = df[df["Salle"] == salle]
            if df_s.empty:
                continue

            nb = len(df_s)
            nb_crit = len(df_s[df_s["Priorite"] == "CRITIQUE"])
            nb_trait = len(df_s[df_s["Traite"]])

            with st.expander(
                f"**{salle}** — {nb} intervention(s)"
                + (f"  |  🔴 {nb_crit} critique(s)" if nb_crit else "")
                + (f"  |  ✅ {nb_trait} traitee(s)" if nb_trait else ""),
                expanded=(nb_crit > 0),
            ):
                for _, row in df_s.sort_values(
                    "Priorite",
                    key=lambda x: x.map({"CRITIQUE": 0, "REMPLACEMENT": 1, "TRAVAUX": 2, "A VERIFIER": 3}),
                ).iterrows():
                    prio  = row["Priorite"]
                    color = COULEURS_PRIO.get(prio, "#888")
                    fond  = FONDS_PRIO.get(prio, "#FFF")
                    trait_badge = '<span style="color:#27AE60;font-weight:700"> ✅ Traite</span>' if row["Traite"] else ""

                    st.markdown(
                        f"""<div style="background:{fond};border-left:4px solid {color};
                        padding:6px 12px;border-radius:4px;margin:4px 0;font-size:0.9rem">
                        <b style="color:{color}">{row['Emoji']} {prio}</b> &nbsp;|&nbsp;
                        <b>{row['Element']}</b> &nbsp;&mdash;&nbsp; {row['Valeur']}
                        {trait_badge}</div>""",
                        unsafe_allow_html=True,
                    )


# ════════════════════════════════════════════════════════════════════════════
# TAB 4 : GESTION DES TRAITEMENTS
# ════════════════════════════════════════════════════════════════════════════
with tab_edition:
    st.markdown('<div class="main-title">✅ Gestion des traitements</div>', unsafe_allow_html=True)
    st.caption("Cochez les interventions realisees. L'etat est conserve pendant la session.")

    df_edit = df_base.copy()
    df_edit["Traite"] = df_edit.index.map(st.session_state.traites)

    col_f1, col_f2 = st.columns(2)
    with col_f1:
        salle_edit = st.selectbox(
            "Filtrer par salle", ["Toutes"] + sorted(df_edit["Salle"].unique().tolist())
        )
    with col_f2:
        prio_edit = st.selectbox("Filtrer par priorite", ["Toutes"] + ORDRE_PRIO)

    df_filtered_edit = df_edit.copy()
    if salle_edit != "Toutes":
        df_filtered_edit = df_filtered_edit[df_filtered_edit["Salle"] == salle_edit]
    if prio_edit != "Toutes":
        df_filtered_edit = df_filtered_edit[df_filtered_edit["Priorite"] == prio_edit]

    # Boutons rapides
    c_tout, c_aucun = st.columns([1, 1])
    with c_tout:
        if st.button("✅ Tout marquer comme traite (vue actuelle)"):
            for idx in df_filtered_edit.index:
                st.session_state.traites[idx] = True
            st.rerun()
    with c_aucun:
        if st.button("↩️ Tout reinitialiser (vue actuelle)"):
            for idx in df_filtered_edit.index:
                st.session_state.traites[idx] = False
            st.rerun()

    st.divider()

    # Liste avec checkboxes
    df_filtered_edit = df_filtered_edit.sort_values(
        "Priorite",
        key=lambda x: x.map({"CRITIQUE": 0, "REMPLACEMENT": 1, "TRAVAUX": 2, "A VERIFIER": 3}),
    )

    for _, row in df_filtered_edit.iterrows():
        idx   = row.name
        prio  = row["Priorite"]
        color = COULEURS_PRIO.get(prio, "#888")
        fond  = FONDS_PRIO.get(prio, "#FFF")

        col_check, col_info = st.columns([1, 11])
        with col_check:
            checked = st.checkbox(
                "", value=st.session_state.traites.get(idx, False), key=f"chk_{idx}"
            )
            st.session_state.traites[idx] = checked
        with col_info:
            st.markdown(
                f"""<div style="background:{fond};border-left:4px solid {color};
                padding:5px 10px;border-radius:4px;margin:2px 0;font-size:0.88rem">
                <b style="color:{color}">{row['Emoji']} {prio}</b> &nbsp;|&nbsp;
                <b>{row['Salle']}</b> &gt; {row['Element']} &nbsp;&mdash;&nbsp; {row['Valeur']}
                </div>""",
                unsafe_allow_html=True,
            )

    # Progression
    total_base  = len(df_base)
    nb_traites  = sum(st.session_state.traites.values())
    pct         = int(nb_traites / total_base * 100) if total_base else 0
    st.divider()
    st.markdown(f"**Progression globale : {nb_traites}/{total_base} traitees ({pct}%)**")
    st.progress(pct / 100)


# ── Boutons d'export dans la sidebar ──────────────────────────────────────────
with st.sidebar:
    df_export = df_base.copy()
    df_export["Traite"] = df_export.index.map(st.session_state.traites)

    try:
        docx_bytes = generer_docx(df_export)
        st.download_button(
            label="📄 Telecharger Word (.docx)",
            data=docx_bytes,
            file_name=f"interventions_lycee_{date.today().strftime('%Y%m%d')}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )
    except ImportError:
        st.warning("python-docx requis : pip install python-docx")

    xlsx_bytes = generer_xlsx(df_export)
    st.download_button(
        label="📊 Telecharger Excel (.xlsx)",
        data=xlsx_bytes,
        file_name=f"interventions_lycee_{date.today().strftime('%Y%m%d')}.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        use_container_width=True,
    )

    st.divider()
    st.caption("pip install streamlit openpyxl python-docx plotly pandas")

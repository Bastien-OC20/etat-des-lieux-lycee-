#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generation des exports DOCX et XLSX."""

import io
from datetime import date

import pandas as pd

from constants import (
    ANNEE_SCOLAIRE,
    COULEURS_PRIO,
    FONDS_PRIO,
    ORDRE_PRIO_MAP,
    hex2rgb,
)


def generer_docx(df: pd.DataFrame, tri_par: str = "Salle") -> bytes:
    """Genere un rapport Word (.docx) des interventions.

    Args:
        df: DataFrame des interventions avec colonne Traite.
        tri_par: cle de regroupement — 'Salle', 'Element' ou 'Priorite'.

    Returns:
        Contenu binaire du fichier .docx.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    def set_cell_bg(cell, hex_color: str) -> None:
        hex_color = hex_color.lstrip("#")
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_para_bg(para, hex_color: str) -> None:
        hex_color = hex_color.lstrip("#")
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def add_border_bottom(para, color: str = "1A3C6E") -> None:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        pPr.append(pBdr)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titre
    p = doc.add_paragraph()
    r = p.add_run("Etat des lieux - Interventions a entreprendre")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(p)

    p2 = doc.add_paragraph()
    date_str = date.today().strftime('%d/%m/%Y')
    p2.add_run(f"Lycee - Annee {ANNEE_SCOLAIRE}  |  Genere le {date_str}")
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # Bilan global
    total_i = len(df)
    non_traites = (
        len(df[~df["Traite"]]) if "Traite" in df.columns else total_i
    )
    p_stat = doc.add_paragraph()
    r_stat = p_stat.add_run(
        f"Total interventions : {total_i}  |  Non traitees : {non_traites}"
    )
    r_stat.bold = True
    r_stat.font.size = Pt(11)
    r_stat.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p_stat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    TRI_CONFIG = {
        "Salle": {
            "col": "Salle",
            "label": "Salle",
            "headers": ["Element", "Valeur originale", "Priorite", "Statut"],
            "row_cols": ["Element", "Valeur", "Priorite"],
            "sort_within": "Priorite",
            "widths": [5, 5, 3.5, 2.5],
        },
        "Element": {
            "col": "Element",
            "label": "Type d'intervention",
            "headers": ["Salle", "Valeur originale", "Priorite", "Statut"],
            "row_cols": ["Salle", "Valeur", "Priorite"],
            "sort_within": "Priorite",
            "widths": [4, 5, 4, 3],
        },
        "Priorite": {
            "col": "Priorite",
            "label": "Urgence",
            "headers": ["Salle", "Element", "Valeur originale", "Statut"],
            "row_cols": ["Salle", "Element", "Valeur"],
            "sort_within": "Salle",
            "widths": [4, 4.5, 5, 2.5],
        },
    }
    cfg = TRI_CONFIG.get(tri_par, TRI_CONFIG["Salle"])

    if tri_par == "Priorite":
        ordered = (
            df.assign(_ordre=df["Priorite"].map(ORDRE_PRIO_MAP).fillna(99))
            .sort_values("_ordre")["Priorite"]
            .unique()
        )
    else:
        ordered = sorted(df[cfg["col"]].unique())

    for groupe_val in ordered:
        df_g = df[df[cfg["col"]] == groupe_val]
        p_hdr = doc.add_paragraph()
        r_hdr = p_hdr.add_run(
            f"  {cfg['label']} : {groupe_val}  ({len(df_g)} intervention(s))"
        )
        r_hdr.bold = True
        r_hdr.font.size = Pt(12)
        bg_hdr = (
            FONDS_PRIO.get(groupe_val, "#1A3C6E").lstrip("#")
            if tri_par == "Priorite"
            else "1A3C6E"
        )
        set_para_bg(p_hdr, bg_hdr)
        if tri_par == "Priorite":
            tc_hdr = COULEURS_PRIO.get(groupe_val, "#FFFFFF")
            r_hdr.font.color.rgb = RGBColor(*hex2rgb(tc_hdr))
        else:
            r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        for cell, label in zip(table.rows[0].cells, cfg["headers"]):
            cell.text = label
            set_cell_bg(cell, "D5E8F0")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for _, row in df_g.sort_values(cfg["sort_within"]).iterrows():
            cells = table.add_row().cells
            for i, col in enumerate(cfg["row_cols"]):
                cells[i].text = str(row[col])
            cells[3].text = "✓ Traite" if row.get("Traite") else "En attente"

            prio_val = row["Priorite"]
            bg = FONDS_PRIO.get(prio_val, "#F5F5F5").lstrip("#")
            tc_hex = COULEURS_PRIO.get(prio_val, "#888888")
            prio_cell_idx = (
                cfg["row_cols"].index("Priorite")
                if "Priorite" in cfg["row_cols"]
                else None
            )
            if prio_cell_idx is not None:
                set_cell_bg(cells[prio_cell_idx], bg)
                r_prio = cells[prio_cell_idx].paragraphs[0].runs[0]
                r_prio.bold = True
                r_prio.font.size = Pt(9)
                r_prio.font.color.rgb = RGBColor(*hex2rgb(tc_hex))
            for c in cells:
                if c.paragraphs[0].runs:
                    c.paragraphs[0].runs[0].font.size = Pt(9)

        for i, w in enumerate(cfg["widths"]):
            for row in table.rows:
                row.cells[i].width = Cm(w)

        doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def generer_xlsx(df: pd.DataFrame, tri_par: str = "Salle") -> bytes:
    """Genere un tableur Excel (.xlsx) des interventions.

    Args:
        df: DataFrame des interventions avec colonne Traite.
        tri_par: cle de tri — 'Salle', 'Element' ou 'Priorite'.

    Returns:
        Contenu binaire du fichier .xlsx.
    """
    from openpyxl.styles import PatternFill, Font, Alignment

    # Styles derives des constantes partagees (pas de duplication)
    FILLS = {
        prio: PatternFill("solid", fgColor=FONDS_PRIO[prio].lstrip("#"))
        for prio in FONDS_PRIO
    }
    FONTS = {
        prio: Font(bold=True, color=COULEURS_PRIO[prio].lstrip("#"))
        for prio in COULEURS_PRIO
    }

    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        export_df = df[
            ["Salle", "Element", "Valeur", "Priorite", "Traite"]
        ].copy()
        export_df["Traite"] = export_df["Traite"].map(
            {True: "Oui", False: "Non"}
        )
        if tri_par == "Priorite":
            export_df["_ordre"] = (
                export_df["Priorite"].map(ORDRE_PRIO_MAP).fillna(99)
            )
            export_df = (
                export_df.sort_values(["_ordre", "Salle", "Element"])
                .drop(columns="_ordre")
            )
        elif tri_par == "Element":
            export_df = export_df.sort_values(["Element", "Priorite", "Salle"])
        else:
            export_df = export_df.sort_values(["Salle", "Priorite", "Element"])
        export_df.to_excel(writer, index=False, sheet_name="Interventions")

        ws = writer.sheets["Interventions"]
        header_fill = PatternFill("solid", fgColor="1A3C6E")
        header_font = Font(bold=True, color="FFFFFF")
        for cell in ws[1]:
            cell.fill = header_fill
            cell.font = header_font
            cell.alignment = Alignment(horizontal="center")

        for row in ws.iter_rows(min_row=2):
            prio = row[3].value
            if prio in FILLS:
                row[3].fill = FILLS[prio]
                row[3].font = FONTS[prio]

        ws.column_dimensions["A"].width = 20
        ws.column_dimensions["B"].width = 30
        ws.column_dimensions["C"].width = 25
        ws.column_dimensions["D"].width = 18
        ws.column_dimensions["E"].width = 10

    return buf.getvalue()

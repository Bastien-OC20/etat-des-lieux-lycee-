#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
interventions_lycee.py
======================
Script en ligne de commande (CLI) pour analyser le fichier Excel
et lister toutes les interventions a entreprendre dans le lycee.

Usage
-----
  python interventions_lycee.py [options]

Options
-------
  -f / --fichier   Chemin vers le fichier xlsx
  -s / --salle     Filtrer sur une salle  ex: --salle 106
  -d / --docx      Exporter le rapport en fichier Word (.docx)
  -o / --sortie    Nom du fichier de sortie (defaut : interventions_lycee.docx)

Exemples
--------
  python interventions_lycee.py
  python interventions_lycee.py --salle Foyer
  python interventions_lycee.py --docx
  python interventions_lycee.py --salle 106 --docx --sortie rapport_106.docx
"""

import sys
import argparse
from pathlib import Path
from datetime import date


# ── Dependances ────────────────────────────────────────────────────────────────
def _check(module, pip):
    try:
        return __import__(module)
    except ImportError:
        print(f"Module manquant : pip install {pip}")
        sys.exit(1)


openpyxl = _check("openpyxl", "openpyxl")

# ── Imports partages depuis constants ─────────────────────────────────────────
from constants import (  # noqa: E402
    ANNEE_SCOLAIRE,
    COULEURS_PRIO,
    EMOJIS_PRIO,
    FONDS_PRIO,
    LABELS_PRIO,
    VALEURS_KO,
    VALEURS_OK,
    hex2rgb,
    normaliser,
)

# ── Configuration ──────────────────────────────────────────────────────────────
_annee_fichier = ANNEE_SCOLAIRE.replace("/", "_")
FICHIER_PAR_DEFAUT = (
    Path(__file__).parent / f"Etat_des_lieux_{_annee_fichier}.xlsx"
)


# ── Chargement ─────────────────────────────────────────────────────────────────
def charger_donnees(chemin: str) -> list[dict]:
    """Charge le fichier Excel et retourne la liste des salles avec leurs valeurs."""
    path = Path(chemin)
    if not path.exists():
        alt = Path(chemin).parent / f"État_des_lieux_{_annee_fichier}.xlsx"
        if alt.exists():
            path = alt
        else:
            print(f"Fichier introuvable : {chemin}")
            sys.exit(1)

    wb = openpyxl.load_workbook(path)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    headers = [str(h).strip() if h else "" for h in rows[1]]
    salles = []
    for row in rows[2:]:
        if row[0] is None:
            continue
        salle: dict = {"nom": str(row[0]).strip(), "colonnes": {}}
        for i, val in enumerate(row[1:], start=1):
            if i < len(headers) and headers[i]:
                salle["colonnes"][headers[i]] = (
                    str(val).strip() if val is not None else ""
                )
        salles.append(salle)
    return salles


# ── Analyse ────────────────────────────────────────────────────────────────────
def analyser(salle: dict) -> list[tuple[str, str, str]]:
    """Retourne [(colonne, priorite, libelle)] pour chaque probleme detecte."""
    interventions = []
    for col, val in salle["colonnes"].items():
        norm = normaliser(val)
        if norm in VALEURS_OK:
            continue
        matched = VALEURS_KO.get(norm)
        if matched is None:
            for motcle, info in VALEURS_KO.items():
                if motcle in norm:
                    matched = info
                    break
        if matched:
            prio, _ = matched
            libelle = LABELS_PRIO.get(prio, prio)
            interventions.append((col, prio, f"{libelle}  ({val})"))
        else:
            prio = "A VERIFIER"
            interventions.append((col, prio, f"{LABELS_PRIO[prio]} : {val}"))
    return interventions


# ── Affichage console ──────────────────────────────────────────────────────────
def afficher_rapport(
    salles: list[dict], filtre: str | None = None
) -> None:
    """Affiche le rapport dans le terminal."""
    L = 72
    print("=" * L)
    print("  ETAT DES LIEUX LYCEE - INTERVENTIONS A ENTREPRENDRE")
    print("=" * L)

    total_s = total_i = 0
    for salle in salles:
        nom = salle["nom"]
        if filtre and filtre.lower() not in nom.lower():
            continue
        interventions = analyser(salle)
        total_s += 1
        if not interventions:
            print(f"\n  OK   Salle {nom} - Aucune intervention requise")
            continue
        total_i += len(interventions)
        print(f"\n  Salle : {nom}  ({len(interventions)} intervention(s))")
        for col, prio, label in interventions:
            emoji = EMOJIS_PRIO.get(prio, "⚪")
            print(f"    {emoji}  {col:<30} {label}")

    print("\n" + "=" * L)
    print(f"  BILAN : {total_s} salle(s)  |  {total_i} intervention(s)")
    print("=" * L)


# ── Export Word (.docx) ────────────────────────────────────────────────────────
def exporter_docx(
    salles: list[dict],
    chemin_sortie: str,
    filtre: str | None = None,
) -> None:
    """Genere un rapport Word (.docx) a partir de la liste des salles."""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        print("Module manquant pour l'export Word : pip install python-docx")
        sys.exit(1)

    def set_cell_bg(cell, hex_color: str) -> None:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        tcPr.append(shd)

    def set_para_bg(para, hex_color: str) -> None:
        pPr = para._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_color)
        pPr.append(shd)

    def add_border_bottom(para, color: str = "CCCCCC") -> None:
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        b = OxmlElement("w:bottom")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "1")
        b.set(qn("w:color"), color)
        pBdr.append(b)
        pPr.append(pBdr)

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Titre
    p = doc.add_paragraph()
    r = p.add_run("Etat des lieux - Interventions a entreprendre")
    r.bold = True
    r.font.size = Pt(18)
    r.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_border_bottom(p, "1A3C6E")

    p2 = doc.add_paragraph()
    p2.add_run(
        f"Lycee - Annee {ANNEE_SCOLAIRE}"
        f"  |  Genere le {date.today().strftime('%d/%m/%Y')}"
    )
    p2.runs[0].font.size = Pt(10)
    p2.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    total_s = total_i = 0
    for salle in salles:
        nom = salle["nom"]
        if filtre and filtre.lower() not in nom.lower():
            continue
        interventions = analyser(salle)
        total_s += 1

        p_hdr = doc.add_paragraph()
        r_hdr = p_hdr.add_run(f"  Salle : {nom}")
        r_hdr.bold = True
        r_hdr.font.size = Pt(12)
        r_hdr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_para_bg(p_hdr, "1A3C6E")

        if not interventions:
            p_ok = doc.add_paragraph()
            r_ok = p_ok.add_run("    OK - Aucune intervention requise")
            r_ok.font.size = Pt(10)
            r_ok.font.color.rgb = RGBColor(0x27, 0xAE, 0x60)
            doc.add_paragraph()
            continue

        total_i += len(interventions)

        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"

        for cell, label in zip(
            table.rows[0].cells,
            ["Element concerne", "Action requise", "Priorite"],
        ):
            cell.text = label
            set_cell_bg(cell, "D5E8F0")
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)

        for col, prio, label in interventions:
            cells = table.add_row().cells
            cells[0].text = col
            cells[0].paragraphs[0].runs[0].font.size = Pt(9)
            cells[1].text = label
            cells[1].paragraphs[0].runs[0].font.size = Pt(9)

            tc_hex = COULEURS_PRIO.get(prio, "#888888").lstrip("#")
            bg_hex = FONDS_PRIO.get(prio, "#F5F5F5").lstrip("#")
            set_cell_bg(cells[1], bg_hex)
            set_cell_bg(cells[2], bg_hex)

            p_prio = cells[2].paragraphs[0]
            r_prio = p_prio.add_run(prio)
            r_prio.bold = True
            r_prio.font.size = Pt(8)
            r_prio.font.color.rgb = RGBColor(*hex2rgb(tc_hex))

        for row in table.rows:
            row.cells[0].width = Cm(6)
            row.cells[1].width = Cm(9)
            row.cells[2].width = Cm(3)

        doc.add_paragraph()

    p_bilan = doc.add_paragraph()
    add_border_bottom(p_bilan, "1A3C6E")
    r_b = p_bilan.add_run(
        f"Bilan : {total_s} salles analysees"
        f"  |  {total_i} interventions identifiees"
    )
    r_b.bold = True
    r_b.font.size = Pt(11)
    r_b.font.color.rgb = RGBColor(0x1A, 0x3C, 0x6E)
    p_bilan.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(chemin_sortie)
    print(f"\nFichier Word enregistre : {chemin_sortie}")


# ── Main ───────────────────────────────────────────────────────────────────────
def main() -> None:
    parser = argparse.ArgumentParser(
        description="Etat des lieux lycee - interventions par salle"
    )
    parser.add_argument(
        "--fichier", "-f",
        default=str(FICHIER_PAR_DEFAUT),
        help="Fichier xlsx source",
    )
    parser.add_argument(
        "--salle", "-s",
        default=None,
        help="Filtrer sur une salle  ex: 106  ou  Foyer",
    )
    parser.add_argument(
        "--docx", "-d",
        action="store_true",
        help="Exporter le rapport en fichier Word (.docx)",
    )
    parser.add_argument(
        "--sortie", "-o",
        default="interventions_lycee.docx",
        help="Nom du fichier Word de sortie",
    )
    args = parser.parse_args()

    salles = charger_donnees(args.fichier)
    afficher_rapport(salles, filtre=args.salle)

    if args.docx:
        p_sortie = Path(args.sortie)
        if not p_sortie.is_absolute():
            p_sortie = Path(__file__).parent / p_sortie
        exporter_docx(salles, str(p_sortie), filtre=args.salle)


if __name__ == "__main__":
    main()
